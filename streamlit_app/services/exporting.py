"""Utilities for generating downloadable artefacts from audit reports."""
from __future__ import annotations

import io
import json
from datetime import datetime
from typing import Any, Dict, Iterable, List, Tuple

import pandas as pd
import streamlit as st
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def _timestamped(name: str, ext: str) -> str:
    slug = name.lower().replace(" ", "-")
    ts = datetime.utcnow().strftime("%Y%m%d-%H%M%S")
    return f"{slug}-{ts}.{ext}"


def export_as_json(report: Dict[str, Any]) -> Tuple[str, bytes]:
    filename = _timestamped(report.get("summary", {}).get("title", "relatorio"), "json")
    payload = json.dumps(report, indent=2, ensure_ascii=False).encode("utf-8")
    return filename, payload


@st.cache_data(show_spinner=False)
def export_documents_dataframe(report: Dict[str, Any]) -> pd.DataFrame:
    documents: Iterable[Dict[str, Any]] = report.get("documents", []) or []
    rows: List[Dict[str, Any]] = []
    for doc in documents:
        meta = doc.get("doc", {}) if isinstance(doc, dict) else {}
        rows.append(
            {
                "Arquivo": meta.get("name"),
                "Tipo": meta.get("kind"),
                "Tamanho": meta.get("size"),
                "Status": doc.get("status"),
                "Inconsistências": len(doc.get("inconsistencies", []) or []),
                "Classificação": (doc.get("classification") or {}).get("operationType"),
            }
        )
    return pd.DataFrame(rows)


def export_as_excel(report: Dict[str, Any]) -> Tuple[str, bytes]:
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        export_documents_dataframe(report).to_excel(writer, sheet_name="Documentos", index=False)
        insights = report.get("incrementalInsights") or []
        if insights:
            insights_df = pd.DataFrame(insights)
            insights_df.to_excel(writer, sheet_name="Insights", index=False)
        metrics = report.get("aggregatedMetrics") or {}
        if metrics:
            metrics_df = pd.DataFrame([metrics])
            metrics_df.to_excel(writer, sheet_name="Métricas", index=False)
    filename = _timestamped(report.get("summary", {}).get("title", "relatorio"), "xlsx")
    return filename, buffer.getvalue()


def export_as_doc(report: Dict[str, Any]) -> Tuple[str, bytes]:
    document = Document()
    summary = report.get("summary", {})
    document.add_heading(summary.get("title", "Relatório Fiscal"), 0)
    document.add_paragraph(summary.get("overview", ""))
    for metric in summary.get("keyMetrics", []):
        document.add_paragraph(f"{metric.get('metric')}: {metric.get('value')} ({metric.get('insight')})")
    document.add_heading("Insights", level=1)
    for insight in report.get("incrementalInsights", []):
        document.add_paragraph(f"- {insight.get('title', 'Insight')}: {insight.get('description', '')}")
    document.add_heading("Histórico do Chat", level=1)
    for message in report.get("chatLog", []):
        document.add_paragraph(f"[{message.get('sender', 'ai').upper()}] {message.get('text', '')}")
    buffer = io.BytesIO()
    document.save(buffer)
    filename = _timestamped(summary.get("title", "relatorio"), "docx")
    return filename, buffer.getvalue()


def export_as_pdf(report: Dict[str, Any]) -> Tuple[str, bytes]:
    buffer = io.BytesIO()
    summary = report.get("summary", {})
    story: List[Any] = []
    styles = getSampleStyleSheet()
    story.append(Paragraph(summary.get("title", "Relatório Fiscal"), styles["Title"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph(summary.get("overview", ""), styles["BodyText"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Principais métricas", styles["Heading2"]))
    for metric in summary.get("keyMetrics", []):
        story.append(Paragraph(f"• <b>{metric.get('metric')}</b>: {metric.get('value')} - {metric.get('insight')}", styles["BodyText"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Insights Incrementais", styles["Heading2"]))
    for insight in report.get("incrementalInsights", []):
        story.append(Paragraph(f"• {insight.get('title', 'Insight')}: {insight.get('description', '')}", styles["BodyText"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Histórico do Chat", styles["Heading2"]))
    for message in report.get("chatLog", []):
        story.append(Paragraph(f"[{message.get('sender', 'ai').upper()}] {message.get('text', '')}", styles["BodyText"]))
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    doc.build(story)
    filename = _timestamped(summary.get("title", "relatorio"), "pdf")
    return filename, buffer.getvalue()
